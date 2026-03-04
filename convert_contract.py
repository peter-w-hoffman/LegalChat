import pandas as pd 
import numpy as np 
import json 
import os
from openai import OpenAI
from docx import Document

client = OpenAI()

data_path = "/Users/carolgao/MIT Dropbox/Carol Gao/common_experience_2025/"

CONTRACT = "2181CV01689"
# Load the Word document
# Contract 1 
word_file = f"{data_path}/contracts/{CONTRACT}_contract.docx"
doc = Document(word_file)
full_text = "\n".join([para.text for para in doc.paragraphs])
response = client.chat.completions.create(
    model="gpt-4o-mini",  
    messages=[
        {"role": "system", "content": "You are a helpful proofreader."},
        {"role": "user", "content": f"Please correct typos in the following t:\n\n{full_text}"}
    ],
)
corrected_text = response.choices[0].message.content
with open(f"{data_path}/contracts/corrected_output_{CONTRACT}.txt", "w", encoding="utf-8") as f:
    f.write(corrected_text)

# Contract 2
#word_file = f"{data_path}/contracts/2277CV00212_contract.docx"
#doc = Document(word_file)
#full_text = "\n".join([para.text for para in doc.paragraphs])
#with open(f"{data_path}/contracts/corrected_output_2277CV00212.txt", "w", encoding="utf-8") as f:
#    f.write(full_text)

